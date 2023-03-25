#一、 引用函数、库等写在最上方
#11111111111111111111111111111111111111111111111111111111111111
import win32com.client as win32
#win32com非常难安装，https://github.com/mhammond/pywin32/releases,下载正确的exe
#如果还出错，再把C:\Users\lilyhcn\AppData\Local\Programs\Python\Python311\Lib\site-packages\pywin32_system32
# 添加到环境变量中的path，感谢chatgpt
import os
from docx import Document
# 打开 Word 应用程序,

# word = win32.gencache.EnsureDispatch("Word.Application")

#有bug，这个会把文本框给弄没了
def replacenewfile(oldpath, replace_dict,newpath):
    document = Document(oldpath)
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    
                    #print(key+"-->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    

    document.save(newpath)
    return newpath

#有bug，这个会把文本框给弄没了
def replacenewfile2(newpath, replace_dict):
    document = Document(newpath)
    for para in document.paragraphs:
        tptext=para.text
        #print(tptext)
        for key, value in replace_dict.items():
            if key in tptext:
                #print(key+"-->"+value)
                para.text = tptext.replace(key, value)
                # print(para.text)
    document.save(newpath)
    return newpath


#有bug，这个只能替换文本框
def checkshapetext(oldpath, replace_dict,newpath=""):
    doc = word.Documents.Open(oldpath)
    # 循环遍历文档中的所有形状
    for shape in doc.Shapes:
        # 如果形状是文本框
        if shape.Type == 17:
            # 获取文本框中的文本
            text = shape.TextFrame.TextRange.Text

            for key, value in replace_dict.items():
                #print(key+"-->"+value+"----text=:"+text)
                if key in text:
                   
                    new_text = text.replace(key, value)
                    shape.TextFrame.TextRange.Text = new_text

    if newpath=="":
        newpath  =oldpath
    # 保存修改后的文档
    doc.SaveAs(newpath)

    # 关闭 Word 文档和应用程序
    doc.Close()
    word.Quit()
    return newpath

#11111111111111111111111111111111111111111111111111111111111111
# ---------------r34.cc制作 excel 的输入输出---------------
import os,base64,time,sys,json,lilyfun # 导入同路径下的函数
prflag = "true"  # 是否打印输出，true输出
inarr,outarr={},{}
lilyfun.tj()

# 二、运行出错时，默认的输入、输出的默认标题行
#2222222222222222222222222222222222222222222222222222222222222
#输入文本
inarr["模板路径"]="D:\\老黄牛小工具\\word模板\\座位牌3字.docx"
inarr["生成文件路径"]="老黄牛.docx"
inarr["姓名"]="老黄牛"
inarr[""]=""
inarr[""]=""
outarr[""] = ""
outarr[""] = ""
outarr[""] = ""
outarr[""] = ""
outarr[""] = ""
#文件标志
fkeyold="模板路径"   #输入变量中，哪个是文件的标记。
fkeynew="生成文件路径"   #输出变量中，哪个是文件的标记。
#2222222222222222222222222222222222222222222222222222222222222
config=lilyfun.readiniconfig()
inarr=lilyfun.updatearrfromini(inarr,config)




def main(fd2={}):
    global inarr,outarr,prflag,fkeyold,fkeynew,mlkey,config
    arr2ret,valarr,errarr = {},{},{}
    wholepath,f64="",""
    errarr=lilyfun.merge(inarr,outarr)  #合并字典
    # ----------------三、初始化读取数据：----------------------
    #读valarr,即读标题及标题对应的值
    # ----------------[1/4]读取fd2，生成base64字符串 json64 -----
    try: #1.1读取fd2，即传入字典
        json64=lilyfun.getfd2(fd2,"json64")
    except:
        #errarr["执行结果"]="读取fd2错误，请检查！"
        errarr = lilyfun.printvalarr(errarr,"读取fd2错误，请检查！",prflag)
        return lilyfun.mboutputarr(fd2,prflag,errarr)
    lilyfun.titlepr("[1/4] fd2 传入成功！：","传入成功",prflag)
    
    # ----------------[2/4] 生成f64,json64解码为jsonarr -----------
    try: #1.2 json64解码为jsonarr
        jsonarr = lilyfun.json64tojsonarr(json64)
        jsoncontentarr =jsonarr["contents"]
        jsoncontentarr=lilyfun.updatearrfromini(jsoncontentarr,config)
        jsonarr["contents"]=jsoncontentarr
        f64=lilyfun.getfd2_f64(fd2,fkeyold,jsonarr)
    except:
        errarr = lilyfun.printvalarr(errarr,"jsonarr解码错误，请检查！",prflag)
        #fd2:函数传过来的值，arr2ret:运行得到的数组，prflag：打印标记
        return lilyfun.mboutputarr(fd2,prflag,errarr)
    lilyfun.titlepr("[2/4] 解码成功 jsonarr：",jsonarr,prflag)

    # ----------------[3/4]按输入区读传过来的值并反馈到字典valarr ------
    try:
        valarr=lilyfun.getvalarr(jsonarr,inarr,outarr,prflag)
    except:
        errarr = lilyfun.printvalarr(errarr,"[读输入]标题行没有需要的值，请确保标题行存在。",prflag)
        return lilyfun.mboutputarr(fd2,prflag,errarr)
    lilyfun.titlepr("","获取到的f64的长度为: " + str(len(f64)),prflag)
    lilyfun.titlepr("[3/4] 检查输入值成功 valarr：",valarr,prflag)
    #这里可能还要再写读文件的
    
    # ----------------[4/4]调用函数并生成arr2ret及f64 -------------------
    try:  # 运行函数,最后要生成arr2ret及f64
        old_filepath=lilyfun.randfile(valarr,fkeyold,"old")
        new_filepath=lilyfun.randfile(jsoncontentarr,fkeynew,"new")
        old_filepath=lilyfun.writefile64(f64,old_filepath)
    except:  # 保存函数出错后的执行结果
        valarr = lilyfun.printvalarr(valarr,"[运行]读写文件错误。",prflag)
        return lilyfun.mboutputarr(fd2,prflag,valarr)


    #3333333333333333333333333333333333333333333333333333
    #txt=mainrun(valarr,old_filepath,new_filepath)
    #inarr 模板路径 生成文件路径 姓名  
    #inarr     
    #fkeyold 模板路径 fkeynew 生成文件路径 
    try:  # 运行函数,最后要生成arr2ret及f64

        #new_filepath =checkshapetext(old_filepath,valarr,new_filepath)
        new_filepath =replacenewfile(old_filepath,valarr,new_filepath)

    except Exception as e:# 保存函数出错后的执行结果
        valarr = lilyfun.printvalarr(valarr,"[运行]调用函数出错，请检查值是否正确。" +"\n"+'错误类型：'+ e.__class__.__name__+"\n"+ '错误明细：'+str(e))
        return lilyfun.mboutputarr(fd2,prflag,valarr)





    except Exception as e:# 保存函数出错后的执行结果
        valarr = lilyfun.printvalarr(valarr,"[运行]调用函数出错，请检查值是否正确。" +"\n"+'错误类型：'+ e.__class__.__name__+"\n"+ '错误明细：'+str(e))
        return lilyfun.mboutputarr(fd2,prflag,valarr)




    #3333333333333333333333333333333333333333333333333333
    
    try:  # 运行函数,最后要生成arr2ret及f64
        f64=lilyfun.readfile2f64(new_filepath)#有新文件就读取
        # newpath = valarr[fkeynew]
        # if f64!="" and fkeynew!="" and fd2=={}:
        #     lilyfun.writefile64(f64,newpath)
        lilyfun.safedel(old_filepath)
        lilyfun.safedel(new_filepath)
        arr2ret["执行结果"]="√"
    except:  # 保存函数出错后的执行结果
        valarr = lilyfun.printvalarr(valarr,"[执行函数后]读写、删除文件错误。",prflag)
        return lilyfun.mboutputarr(fd2,prflag,valarr)
    lilyfun.titlepr("[4/4] 函数执行成功 arr2ret：",arr2ret,prflag)
    

    # ----------------五、写入文件，并返回字典 -------------------
    try:  # 写入文件
        if fd2=={} and fkeynew !="":
            excelfolder=lilyfun.safegetkey(jsonarr,"excelpath")
            raltiveapth=jsonarr["contents"][fkeynew]
            wholepath=lilyfun.getwholepath(raltiveapth,excelfolder)
        lilyfun.titlepr("执行、输出成功。","",prflag)
        #这是最关键的返回函数，并写入文件
        #print(wholepath)
        return lilyfun.mboutputarr(fd2,prflag,arr2ret,f64,wholepath,"key")
    except:  # 保存函数出错后的执行结果
        valarr = lilyfun.printvalarr(valarr,"写入文件出错，请检查值是否正确。",prflag)
        lilyfun.titlepr("最后写入文件出错，请检查值是否正确。","",prflag)       
        return lilyfun.mboutputarr(fd2,prflag,valarr)



if __name__ == '__main__':
    main()

    #fd2的内容
    # "json64": json文件  上传过来的json
    # "f64": f64   上传的文件的base64（只能一个文件）
    # "fkeyold": fkeyold  上传时的标题行（只能一个文件）
    # "fkeynew": fkeynew  返回时的标题行（只能一个文件）
        #fd2,prflag="true",arr2ret,f64="",fkeynew="",keyflag="all"
        #fd2:传过来的值，prflag：打印标记,keyflag:excel中的是否全部输出
        #arr2ret:运行得到的字典，f64:反馈文件的base64,fkeynew：输出值
        #return lilyfun.mboutputarr(fd2,prflag,errarr)


